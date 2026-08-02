"""
CV Engine V4 — Professional colors, dramatic layout variety
"""
import random
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml.etree import SubElement

def _rgb(h): return RGBColor(int(h[:2],16),int(h[2:4],16),int(h[4:6],16))
def _border(p,color,sz="6",style="single"):
    pPr=p._p.get_or_add_pPr();pBdr=SubElement(pPr,qn('w:pBdr'))
    b=SubElement(pBdr,qn('w:bottom'));b.set(qn('w:val'),style);b.set(qn('w:sz'),sz);b.set(qn('w:color'),color);b.set(qn('w:space'),'4')
def _top_border(p,color,sz="6"):
    pPr=p._p.get_or_add_pPr();pBdr=SubElement(pPr,qn('w:pBdr'))
    b=SubElement(pBdr,qn('w:top'));b.set(qn('w:val'),'single');b.set(qn('w:sz'),sz);b.set(qn('w:color'),color);b.set(qn('w:space'),'4')
def _shading_run(r,color):
    rPr=r._r.get_or_add_rPr();shd=SubElement(rPr,qn('w:shd'));shd.set(qn('w:val'),'clear');shd.set(qn('w:fill'),color)
def _shading_cell(cell,color):
    tc=cell._tc;tcPr=tc.get_or_add_tcPr();shd=SubElement(tcPr,qn('w:shd'));shd.set(qn('w:fill'),color);shd.set(qn('w:val'),'clear')
def _no_borders(table):
    tbl=table._tbl;tblPr=tbl.tblPr if tbl.tblPr is not None else SubElement(tbl,qn('w:tblPr'))
    borders=SubElement(tblPr,qn('w:tblBorders'))
    for edge in['top','left','bottom','right','insideH','insideV']:
        e=SubElement(borders,qn(f'w:{edge}'));e.set(qn('w:val'),'none');e.set(qn('w:sz'),'0')
def _tab_right(p):
    pPr=p._p.get_or_add_pPr();tabs=SubElement(pPr,qn('w:tabs'));tab=SubElement(tabs,qn('w:tab'));tab.set(qn('w:val'),'right');tab.set(qn('w:pos'),'9360')

# 52 High-contrast curated professional palettes
COLORS = [
    {"ac":"2D3436","lt":"F5F6FA","nm":"Graphite"},
    {"ac":"0C3547","lt":"EBF5FB","nm":"Deep Sea"},
    {"ac":"1E3A5F","lt":"EAF0F7","nm":"Corporate Blue"},
    {"ac":"2C3E50","lt":"ECF0F1","nm":"Slate"},
    {"ac":"1A472A","lt":"EDF5F0","nm":"Pine"},
    {"ac":"3C1518","lt":"F5EEEF","nm":"Wine"},
    {"ac":"2B2D42","lt":"EEEEF2","nm":"Midnight"},
    {"ac":"344E41","lt":"EDF2EF","nm":"Forest"},
    {"ac":"3D405B","lt":"EDEDF2","nm":"Storm"},
    {"ac":"283618","lt":"F0F2EB","nm":"Olive"},
    {"ac":"4A4238","lt":"F2F0ED","nm":"Espresso"},
    {"ac":"1B4332","lt":"EAF4EE","nm":"Emerald"},
    {"ac":"003049","lt":"E6EEF4","nm":"Navy"},
    {"ac":"3A0CA3","lt":"EDEBF7","nm":"Royal"},
    {"ac":"495057","lt":"F1F2F3","nm":"Ash"},
    {"ac":"5C374C","lt":"F2ECF0","nm":"Plum"},
    {"ac":"3E1F47","lt":"F0EBF3","nm":"Grape"},
    {"ac":"14213D","lt":"E8EBF2","nm":"Oxford"},
    {"ac":"4A5859","lt":"EFF1F1","nm":"Steel"},
    {"ac":"2F4858","lt":"ECF0F3","nm":"Petrol"},
    {"ac":"004D40","lt":"E0F2F1","nm":"Deep Teal"},
    {"ac":"5D4037","lt":"EFEBE9","nm":"Warm Copper"},
    {"ac":"212121","lt":"F5F5F5","nm":"Charcoal"},
    {"ac":"1A237E","lt":"E8EAF6","nm":"Navy Slate"},
    {"ac":"33691E","lt":"F1F8E9","nm":"Sage"},
    {"ac":"4E342E","lt":"EFEBE9","nm":"Mocha"},
    {"ac":"006064","lt":"E0F7FA","nm":"Oceanic"},
    {"ac":"4A148C","lt":"F3E5F5","nm":"Deep Purple"},
    {"ac":"880E4F","lt":"FCE4EC","nm":"Rosewood"},
    {"ac":"E65100","lt":"FFF3E0","nm":"Amber Slate"},
    {"ac":"1B5E20","lt":"E8F5E9","nm":"Deep Moss"},
    {"ac":"263238","lt":"ECEFF1","nm":"Midnight Slate"},
    {"ac":"3E2723","lt":"EFEBE9","nm":"Dark Bronze"},
    {"ac":"0D47A1","lt":"E8EAF6","nm":"Cobalt"},
    {"ac":"3E2723","lt":"F5F2F0","nm":"Clay Accent"},
    {"ac":"0A192F","lt":"F0F4F8","nm":"Midnight Blue"},
    {"ac":"4A0E17","lt":"FDF2F4","nm":"Burgundy"},
    {"ac":"143622","lt":"F0F7F2","nm":"Forest Pine"},
    {"ac":"1C2541","lt":"F2F4F8","nm":"Slate Charcoal"},
    {"ac":"6F1D1B","lt":"FDF5F5","nm":"Warm Terra"},
    {"ac":"2E0249","lt":"F9F4FC","nm":"Deep Violet"},
    {"ac":"03045E","lt":"EDF2FB","nm":"Ocean Navy"},
    {"ac":"2B1E16","lt":"F8F6F4","nm":"Espresso Dark"},
    {"ac":"2B3A42","lt":"F0F4F5","nm":"Nordic Steel"},
    {"ac":"064E3B","lt":"ECFDF5","nm":"Emerald Deep"},
    {"ac":"1E1B4B","lt":"EEF2FF","nm":"Indigo Slate"},
    {"ac":"451A03","lt":"FFFBEB","nm":"Bronze Roast"},
    {"ac":"881337","lt":"FFF1F2","nm":"Rose Maroon"},
    {"ac":"134E4A","lt":"F0FDFA","nm":"Teal Dark"},
    {"ac":"172554","lt":"EFF6FF","nm":"Cobalt Royal"},
    {"ac":"18181B","lt":"F4F4F5","nm":"Graphite Dark"}
]

FONT_PAIRS = [
    {"h": "Calibri", "b": "Calibri"},
    {"h": "Arial", "b": "Calibri"},
    {"h": "Georgia", "b": "Garamond"},
    {"h": "Trebuchet MS", "b": "Tahoma"},
    {"h": "Century Gothic", "b": "Calibri"},
    {"h": "Palatino Linotype", "b": "Book Antiqua"},
    {"h": "Cambria", "b": "Georgia"},
    {"h": "Verdana", "b": "Arial"},
    {"h": "Garamond", "b": "Georgia"},
    {"h": "Segoe UI", "b": "Calibri"},
    {"h": "Bookman Old Style", "b": "Garamond"},
    {"h": "Helvetica", "b": "Calibri"},
    {"h": "Corbel", "b": "Calibri"},
    {"h": "Gill Sans MT", "b": "Tahoma"},
    {"h": "Georgia", "b": "Calibri"}
]

BULLETS = ["•","–","▹","›","◦","·","▪","—","→","✧","🔹","❖"]


class CVBuilder:
    def __init__(self):
        self.pal = random.choice(COLORS)
        self.font_pair = random.choice(FONT_PAIRS)
        self.head_font = self.font_pair["h"]
        self.body_font = self.font_pair["b"]
        self.bullet = random.choice(BULLETS)
        self.name_style = random.choice(["center","left","spaced","banner_dark","banner_light","two_tone","minimal","uppercase_line","split_right","underlined_name","top_double_line","left_border_stripe","accent_box","split_header_pills","sidebar_contact_block","modern_compact_header"])
        self.head_style = random.choice(["line","thick_line","block_dark","block_light","bar_left","caps_only","dotted","top_accent","double_underline","full_shaded_strip","right_aligned_accent","boxed_header","pill_badge_title","left_vertical_accent"])
        self.skill_layout = random.choice(["two_col","inline_dots","tag_list","simple_list","category_blocks","bold_pill_tags","three_col_grid","shaded_table_cells"])
        self.body_sz = random.choice([9.5,10,10.5])
        self.head_sz = random.choice([10.5,11,12])
        self.name_sz = random.choice([18,20,22,24,26])
        self.margin = random.choice([1.6,1.8,2.0,2.2])
        self.contact_style = random.choice(["pipe","dot","dash","newline"])
        
        self.doc = Document()
        for s in self.doc.sections:
            s.top_margin=Cm(self.margin-0.4);s.bottom_margin=Cm(self.margin-0.4)
            s.left_margin=Cm(self.margin);s.right_margin=Cm(self.margin)
        
        self.ac = self.pal["ac"]
        self.lt = self.pal["lt"]
    
    def _r(self,p,text,sz=None,bold=False,italic=False,color=None,is_head=False):
        r=p.add_run(text);r.font.size=Pt(sz or self.body_sz)
        r.font.name = self.head_font if is_head else self.body_font
        r.bold=bold;r.italic=italic
        if color:r.font.color.rgb=_rgb(color) if isinstance(color,str) else color
        return r

    def add_name(self, name, email, phone, city):
        sep = {"pipe":" | ","dot":" · ","dash":" — ","newline":"\n"}[self.contact_style]
        contact = f"{email}{sep}{phone}{sep}{city}"
        
        width_cm = 21.0 - (2 * self.margin)
        
        if self.name_style == "banner_dark":
            t=self.doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;_no_borders(t)
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm)
            cell=t.rows[0].cells[0];cell.width = Cm(width_cm);_shading_cell(cell,self.ac)
            p=cell.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_before=Pt(18);p.space_after=Pt(6)
            self._r(p,name.upper(),self.name_sz,bold=True,color="FFFFFF",is_head=True)
            p2=cell.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER;p2.space_after=Pt(12)
            self._r(p2,f"{email}{sep}{phone}{sep}{city}",9,color="CCCCCC")
        
        elif self.name_style == "banner_light":
            t=self.doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;_no_borders(t)
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm)
            cell=t.rows[0].cells[0];cell.width = Cm(width_cm);_shading_cell(cell,self.lt)
            p=cell.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_before=Pt(16);p.space_after=Pt(6)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            p2=cell.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER;p2.space_after=Pt(12)
            self._r(p2,f"{email}{sep}{phone}{sep}{city}",9,color="666666")
        
        elif self.name_style == "two_tone":
            parts=name.upper().split()
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(3)
            if len(parts)>=2:
                self._r(p,parts[0]+" ",self.name_sz,bold=True,color=self.ac,is_head=True)
                self._r(p," ".join(parts[1:]),self.name_sz,bold=False,color="444444",is_head=True)
            else:
                self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(10)
            self._r(p,f"{email}{sep}{phone}{sep}{city}",9,color="888888")
        
        elif self.name_style == "left":
            p=self.doc.add_paragraph();p.space_after=Pt(1)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            p=self.doc.add_paragraph();p.space_after=Pt(10)
            _border(p,self.ac,"4")
            self._r(p,f"{email}{sep}{phone}{sep}{city}",9,color="888888")
        
        elif self.name_style == "spaced":
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(3)
            self._r(p,"    ".join(name.upper().split()),self.name_sz,bold=True,color=self.ac,is_head=True)
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(10)
            _border(p,self.ac,"6")
            self._r(p,f"{email}{sep}{phone}{sep}{city}",9,color="888888")
        
        elif self.name_style == "minimal":
            p=self.doc.add_paragraph();p.space_after=Pt(1)
            self._r(p,name,self.name_sz,bold=True,color="222222",is_head=True)
            p=self.doc.add_paragraph();p.space_after=Pt(10)
            self._r(p,f"{email}{sep}{phone}{sep}{city}",9,color="999999")
        
        elif self.name_style == "uppercase_line":
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(3)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            # Thin accent line
            p2=self.doc.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER;p2.space_after=Pt(3)
            self._r(p2,"━" * 20,8,color=self.ac)
            p3=self.doc.add_paragraph();p3.alignment=WD_ALIGN_PARAGRAPH.CENTER;p3.space_after=Pt(10)
            self._r(p3,f"{email}{sep}{phone}{sep}{city}",9,color="888888")
        
        elif self.name_style == "underlined_name":
            p=self.doc.add_paragraph();p.space_after=Pt(3)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            _border(p,self.ac,"12")
            p2=self.doc.add_paragraph();p2.space_after=Pt(10)
            self._r(p2,f"{email}{sep}{phone}{sep}{city}",9,color="666666")

        elif self.name_style == "top_double_line":
            p_top=self.doc.add_paragraph();p_top.alignment=WD_ALIGN_PARAGRAPH.CENTER;p_top.space_after=Pt(2)
            _top_border(p_top,self.ac,"8")
            self._r(p_top,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            _border(p_top,self.ac,"8")
            p2=self.doc.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER;p2.space_after=Pt(10)
            self._r(p2,f"{email}{sep}{phone}{sep}{city}",9,color="666666")

        elif self.name_style == "accent_box":
            t=self.doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;_no_borders(t)
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm)
            cell=t.rows[0].cells[0];cell.width = Cm(width_cm);_shading_cell(cell,self.lt)
            p=cell.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.LEFT;p.space_before=Pt(12);p.space_after=Pt(4)
            self._r(p,f"  {name.upper()}",self.name_sz,bold=True,color=self.ac,is_head=True)
            p2=cell.add_paragraph();p2.alignment=WD_ALIGN_PARAGRAPH.LEFT;p2.space_after=Pt(10)
            self._r(p2,f"  {email}  ·  {phone}  ·  {city}",9,color="444444")

        elif self.name_style == "split_header_pills":
            p=self.doc.add_paragraph();p.space_after=Pt(3)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            p2=self.doc.add_paragraph();p2.space_after=Pt(10)
            for item in [email, phone, city]:
                r=self._r(p2,f" {item} ",9,color=self.ac)
                _shading_run(r,self.lt)
                self._r(p2,"  ",9)

        elif self.name_style == "sidebar_contact_block":
            p=self.doc.add_paragraph();p.space_after=Pt(2)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            p2=self.doc.add_paragraph();p2.space_after=Pt(10)
            self._r(p2,f"📧 {email}   📱 {phone}   📍 {city}",9,color="555555")

        elif self.name_style == "modern_compact_header":
            p=self.doc.add_paragraph();p.space_after=Pt(1)
            self._r(p,name,self.name_sz,bold=True,color=self.ac,is_head=True)
            p2=self.doc.add_paragraph();p2.space_after=Pt(10)
            _border(p2,self.ac,"6")
            self._r(p2,f"PORTFOLIO CV  |  {email}  |  {phone}  |  {city}",9,color="666666")

        elif self.name_style == "left_border_stripe":
            p=self.doc.add_paragraph();p.space_after=Pt(1)
            self._r(p,f"▌ {name.upper()}",self.name_sz,bold=True,color=self.ac,is_head=True)
            p2=self.doc.add_paragraph();p2.space_after=Pt(10)
            self._r(p2,f"   {email}  ·  {phone}  ·  {city}",9,color="666666")

        elif self.name_style == "split_right":
            t=self.doc.add_table(rows=1,cols=2);t.alignment=WD_TABLE_ALIGNMENT.CENTER;_no_borders(t)
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm * 0.6)
            t.columns[1].width = Cm(width_cm * 0.4)
            t.rows[0].cells[0].width = Cm(width_cm * 0.6)
            t.rows[0].cells[1].width = Cm(width_cm * 0.4)
            
            c0 = t.rows[0].cells[0].paragraphs[0]
            c0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._r(c0, name.upper(), self.name_sz, bold=True, color=self.ac, is_head=True)
            
            c1 = t.rows[0].cells[1].paragraphs[0]
            c1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            p_email = c1
            p_email.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_email.space_after = Pt(2)
            self._r(p_email, email, 9, color="666666")
            
            p_phone = t.rows[0].cells[1].add_paragraph()
            p_phone.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_phone.space_after = Pt(2)
            self._r(p_phone, phone, 9, color="666666")
            
            p_city = t.rows[0].cells[1].add_paragraph()
            p_city.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p_city.space_after = Pt(10)
            self._r(p_city, city, 9, color="666666")
            
        else:  # center
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(3)
            self._r(p,name.upper(),self.name_sz,bold=True,color=self.ac,is_head=True)
            p=self.doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.space_after=Pt(10)
            self._r(p,f"{email}{sep}{phone}{sep}{city}",9,color="888888")
    
    def add_section(self, title):
        p=self.doc.add_paragraph();p.space_before=Pt(13);p.space_after=Pt(5)
        
        if self.head_style == "block_dark":
            r=self._r(p,f"  {title.upper()}  ",self.head_sz,bold=True,color="FFFFFF",is_head=True)
            _shading_run(r,self.ac)
        elif self.head_style == "block_light":
            r=self._r(p,f"  {title.upper()}  ",self.head_sz,bold=True,color=self.ac,is_head=True)
            _shading_run(r,self.lt)
        elif self.head_style == "bar_left":
            self._r(p,f"▎ {title.upper()}",self.head_sz,bold=True,color=self.ac,is_head=True)
        elif self.head_style == "caps_only":
            self._r(p,title.upper(),self.head_sz+1,bold=True,color=self.ac,is_head=True)
        elif self.head_style == "dotted":
            self._r(p,title.upper(),self.head_sz,bold=True,color=self.ac,is_head=True)
            _border(p,self.ac,"4","dotted")
        elif self.head_style == "top_accent":
            _top_border(p,self.ac,"10")
            self._r(p,title.upper(),self.head_sz,bold=True,color=self.ac,is_head=True)
        elif self.head_style == "thick_line":
            self._r(p,title.upper(),self.head_sz,bold=True,color=self.ac,is_head=True)
            _border(p,self.ac,"12")
        elif self.head_style == "double_underline":
            self._r(p,title.upper(),self.head_sz,bold=True,color=self.ac,is_head=True)
            _border(p,self.ac,"6","double")
        elif self.head_style == "full_shaded_strip":
            width_cm = 21.0 - (2 * self.margin)
            t=self.doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;_no_borders(t)
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm)
            cell=t.rows[0].cells[0];cell.width = Cm(width_cm);_shading_cell(cell,self.ac)
            p_cell=cell.paragraphs[0];p_cell.space_before=Pt(3);p_cell.space_after=Pt(3)
            self._r(p_cell,f"  {title.upper()}",self.head_sz,bold=True,color="FFFFFF",is_head=True)
        elif self.head_style == "right_aligned_accent":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._r(p,f"{title.upper()} ━",self.head_sz,bold=True,color=self.ac,is_head=True)
        elif self.head_style == "boxed_header":
            width_cm = 21.0 - (2 * self.margin)
            t=self.doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm)
            cell=t.rows[0].cells[0];cell.width = Cm(width_cm);_shading_cell(cell,self.lt)
            p_cell=cell.paragraphs[0];p_cell.alignment=WD_ALIGN_PARAGRAPH.CENTER;p_cell.space_before=Pt(3);p_cell.space_after=Pt(3)
            self._r(p_cell,title.upper(),self.head_sz,bold=True,color=self.ac,is_head=True)
        elif self.head_style == "pill_badge_title":
            r=self._r(p,f"  {title.upper()}  ",self.head_sz,bold=True,color=self.ac,is_head=True)
            _shading_run(r,self.lt)
        elif self.head_style == "left_vertical_accent":
            self._r(p,f"█  {title.upper()}",self.head_sz,bold=True,color=self.ac,is_head=True)
        else:  # line
            self._r(p,title.upper(),self.head_sz,bold=True,color=self.ac,is_head=True)
            _border(p,self.ac,"4")
    
    def add_entry(self, left, right):
        p=self.doc.add_paragraph();p.space_before=Pt(6);p.space_after=Pt(1)
        self._r(p,left,self.body_sz+0.5,bold=True,color="222222")
        p.add_run("\t")
        self._r(p,right,self.body_sz,italic=True,color="888888")
        _tab_right(p)
    
    def add_sub(self,text):
        p=self.doc.add_paragraph();p.space_after=Pt(1)
        self._r(p,text,self.body_sz,italic=True,color="777777")
    
    def add_bullet(self,text):
        p=self.doc.add_paragraph();p.space_after=Pt(1)
        self._r(p,f"  {self.bullet} {text}",self.body_sz,color="444444")
    
    def add_text(self,text):
        p=self.doc.add_paragraph();p.space_after=Pt(3)
        self._r(p,text,self.body_sz,color="444444")
    
    def add_skills(self, left, right):
        if self.skill_layout == "two_col":
            width_cm = 21.0 - (2 * self.margin)
            t=self.doc.add_table(rows=max(len(left),len(right)),cols=2);_no_borders(t)
            t.allow_autofit = False
            t.columns[0].width = Cm(width_cm * 0.5)
            t.columns[1].width = Cm(width_cm * 0.5)
            for i in range(max(len(left),len(right))):
                t.rows[i].cells[0].width = Cm(width_cm * 0.5)
                t.rows[i].cells[1].width = Cm(width_cm * 0.5)
                if i<len(left):
                    c=t.rows[i].cells[0].paragraphs[0];c.clear();self._r(c,f"  {self.bullet} {left[i]}",self.body_sz,color="444444")
                if i<len(right):
                    c=t.rows[i].cells[1].paragraphs[0];c.clear();self._r(c,f"  {self.bullet} {right[i]}",self.body_sz,color="444444")
        elif self.skill_layout == "three_col_grid":
            width_cm = 21.0 - (2 * self.margin)
            all_s = left + right
            num_rows = (len(all_s) + 2) // 3
            t=self.doc.add_table(rows=num_rows, cols=3); _no_borders(t)
            t.allow_autofit = False
            c_w = Cm(width_cm / 3.0)
            for r_idx in range(num_rows):
                for c_idx in range(3):
                    s_idx = r_idx * 3 + c_idx
                    cell = t.rows[r_idx].cells[c_idx]
                    cell.width = c_w
                    if s_idx < len(all_s):
                        p = cell.paragraphs[0]; p.clear()
                        self._r(p, f"  {self.bullet} {all_s[s_idx]}", self.body_sz, color="444444")
        elif self.skill_layout == "shaded_table_cells":
            width_cm = 21.0 - (2 * self.margin)
            all_s = left + right
            num_rows = (len(all_s) + 1) // 2
            t=self.doc.add_table(rows=num_rows, cols=2)
            t.allow_autofit = False
            c_w = Cm(width_cm * 0.5)
            for r_idx in range(num_rows):
                for c_idx in range(2):
                    s_idx = r_idx * 2 + c_idx
                    cell = t.rows[r_idx].cells[c_idx]
                    cell.width = c_w
                    if s_idx < len(all_s):
                        _shading_cell(cell, self.lt)
                        p = cell.paragraphs[0]; p.clear()
                        self._r(p, f"  {all_s[s_idx]}", self.body_sz, bold=True, color=self.ac)
        elif self.skill_layout == "category_blocks":
            p1 = self.doc.add_paragraph(); p1.space_after = Pt(2)
            self._r(p1, "Technical Competencies: ", self.body_sz, bold=True, color=self.ac)
            self._r(p1, ", ".join(left), self.body_sz, color="444444")
            p2 = self.doc.add_paragraph(); p2.space_after = Pt(4)
            self._r(p2, "Professional Assets: ", self.body_sz, bold=True, color=self.ac)
            self._r(p2, ", ".join(right), self.body_sz, color="444444")
        elif self.skill_layout == "bold_pill_tags":
            all_s = left + right; p = self.doc.add_paragraph(); p.space_after = Pt(3)
            for i, s in enumerate(all_s):
                r = self._r(p, f"  {s}  ", self.body_sz, bold=True, color=self.ac)
                _shading_run(r, self.lt)
                if i < len(all_s) - 1: self._r(p, "  ", self.body_sz)
        elif self.skill_layout == "inline_dots":
            all_s=left+right;p=self.doc.add_paragraph();p.space_after=Pt(3)
            for i,s in enumerate(all_s):
                self._r(p,s,self.body_sz,color=self.ac)
                if i<len(all_s)-1:self._r(p,"  ·  ",self.body_sz,color="CCCCCC")
        elif self.skill_layout == "tag_list":
            all_s=left+right;p=self.doc.add_paragraph();p.space_after=Pt(3)
            for i,s in enumerate(all_s):
                r=self._r(p,f" {s} ",self.body_sz,color=self.ac)
                _shading_run(r,self.lt)
                if i<len(all_s)-1:self._r(p,"  ",self.body_sz)
        else:  # simple_list
            for s in left+right:self.add_bullet(s)
    
    def get_id(self):
        return f"{self.pal['nm']}_{self.name_style}_{self.head_style}_{self.skill_layout}"
    
    def save(self,path):self.doc.save(str(path))


def build_cv(data, output_path):
    cv = CVBuilder()
    d = data
    name = " ".join(d["name"].split())
    
    cv.add_name(name, d["email"], d["phone"], d["city"])
    
    # 8 distinct section orders for structural layout diversity across generations
    section_order = random.choice([
        ["Profile", "Education", "Experience", "Skills", "Languages", "Interests"],
        ["Profile", "Experience", "Education", "Skills", "Languages", "Interests"],
        ["Profile", "Skills", "Education", "Experience", "Languages", "Interests"],
        ["Profile", "Education", "Skills", "Experience", "Languages", "Interests"],
        ["Profile", "Experience", "Skills", "Education", "Languages", "Interests"],
        ["Profile", "Skills", "Experience", "Education", "Languages", "Interests"],
        ["Profile", "Education", "Experience", "Languages", "Skills", "Interests"],
        ["Profile", "Skills", "Education", "Languages", "Experience", "Interests"]
    ])
    
    def render_profile():
        cv.add_section("Profile")
        cv.add_text(d["profile_summary"])

    def render_education():
        edu = d["education"]
        cv.add_section("Education")
        cv.add_entry(edu["degree"], d.get("edu_years","2023 – 2026"))
        cv.add_sub(edu["university"])
        cv.add_bullet(f"Specialization: {edu['specialization']}")
        cv.add_bullet(f"Capstone Project: {edu['project']}")
        coursework = edu.get("coursework", "Solar PV System Sizing, Inverter Electronics, Financial Modeling & Project Management.")
        academic_standing = edu.get("academic_standing", "Ranked in the top 10% of the class. Green Pathways scholarship recipient.")
        cv.add_bullet(f"Core Coursework: {coursework}")
        cv.add_bullet(f"Academic Standing: {academic_standing}")
        
        cv.add_entry(f"Secondary — {edu['secondary_spec']}", d.get("sec_year","2023"))
        cv.add_sub(edu["secondary_school"])
        cv.add_bullet(f"Final Academic Honors: {edu['grade']}")

    def render_experience():
        exp = d["experience"]
        cv.add_section("Experience")
        cv.add_entry(exp["volunteer_org"], d.get("vol_dates","2025 – Present"))
        vol_role = exp.get("volunteer_role", "Volunteer Environmental Educator")
        cv.add_sub(vol_role)
        for b in exp["volunteer_bullets"]: cv.add_bullet(b)
        
        cv.add_entry(exp["placement_org"], d.get("stage_date","June 2024"))
        placement_role = exp.get("placement_role", "Renewable Energy & Sustainability Intern")
        cv.add_sub(placement_role)
        for b in exp["placement_bullets"]: cv.add_bullet(b)

    def render_skills():
        cv.add_section("Skills")
        cv.add_skills(d["cv_skills"]["left"], d["cv_skills"]["right"])

    def render_languages():
        cv.add_section("Languages")
        for l in d["languages"]: cv.add_bullet(l)

    def render_interests():
        cv.add_section("Interests")
        cv.add_text(d["interests"])

    section_map = {
        "Profile": render_profile,
        "Education": render_education,
        "Experience": render_experience,
        "Skills": render_skills,
        "Languages": render_languages,
        "Interests": render_interests
    }

    for sec in section_order:
        section_map[sec]()

    cv.save(output_path)
    return cv.get_id()
